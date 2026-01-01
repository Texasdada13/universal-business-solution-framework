"""
Document Builder - Universal Business Solution Framework

Professional document generation for Word and PDF formats.
Supports templates, styling, and structured document creation.

Usage:
```python
from patterns.export import DocumentBuilder, DocumentStyle

# Create a Word document
doc = DocumentBuilder(title="Project Proposal")
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph("This proposal outlines...")
doc.add_table(data, headers=['Name', 'Value', 'Status'])
doc.add_bullet_list(['Point 1', 'Point 2', 'Point 3'])
doc.save("proposal.docx")

# Create with custom styling
doc = DocumentBuilder(title="Report", style=DocumentStyle.CORPORATE)
```
"""

from dataclasses import dataclass, field
from enum import Enum
from typing import Dict, List, Optional, Any, Union, Tuple
from pathlib import Path
from datetime import datetime
import json


class DocumentStyle(Enum):
    """Pre-defined document styles"""
    DEFAULT = "default"
    CORPORATE = "corporate"
    MINIMAL = "minimal"
    EXECUTIVE = "executive"
    TECHNICAL = "technical"


@dataclass
class StyleConfig:
    """Document style configuration"""
    title_font: str = "Calibri"
    title_size: int = 24
    title_color: str = "1F497D"  # Dark blue

    heading1_font: str = "Calibri"
    heading1_size: int = 16
    heading1_color: str = "1F497D"

    heading2_font: str = "Calibri"
    heading2_size: int = 14
    heading2_color: str = "4472C4"

    body_font: str = "Calibri"
    body_size: int = 11

    table_header_color: str = "4472C4"
    table_header_text_color: str = "FFFFFF"
    table_row_alt_color: str = "F2F2F2"

    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.25
    margin_right: float = 1.25

    @classmethod
    def from_style(cls, style: DocumentStyle) -> 'StyleConfig':
        """Create config from style preset"""
        if style == DocumentStyle.CORPORATE:
            return cls(
                title_color="003366",
                heading1_color="003366",
                heading2_color="0066CC",
                table_header_color="003366"
            )
        elif style == DocumentStyle.MINIMAL:
            return cls(
                title_size=20,
                heading1_size=14,
                heading2_size=12,
                title_color="000000",
                heading1_color="333333",
                heading2_color="666666",
                table_header_color="666666"
            )
        elif style == DocumentStyle.EXECUTIVE:
            return cls(
                title_font="Georgia",
                heading1_font="Georgia",
                heading2_font="Georgia",
                body_font="Georgia",
                title_color="1a1a1a",
                heading1_color="1a1a1a",
                margin_top=1.5,
                margin_bottom=1.5
            )
        elif style == DocumentStyle.TECHNICAL:
            return cls(
                title_font="Consolas",
                body_font="Consolas",
                body_size=10,
                title_color="2E7D32",
                heading1_color="2E7D32",
                table_header_color="2E7D32"
            )
        else:
            return cls()


class DocumentBuilder:
    """
    Professional document builder for Word documents.

    Creates well-formatted documents with consistent styling.
    Falls back to Markdown if python-docx is not available.

    Example:
    ```python
    doc = DocumentBuilder(title="Annual Report")

    doc.add_heading("Financial Summary", level=1)
    doc.add_paragraph("Revenue increased by 15% year-over-year.")

    doc.add_heading("Revenue by Quarter", level=2)
    doc.add_table(
        data=[
            {'Quarter': 'Q1', 'Revenue': '$1.2M'},
            {'Quarter': 'Q2', 'Revenue': '$1.4M'},
        ],
        headers=['Quarter', 'Revenue']
    )

    doc.add_heading("Key Achievements", level=2)
    doc.add_bullet_list([
        'Launched 3 new products',
        'Expanded to 5 new markets',
        'Reduced costs by 10%'
    ])

    doc.save("report.docx")
    ```
    """

    def __init__(
        self,
        title: str = "",
        subtitle: str = "",
        author: str = "",
        style: DocumentStyle = DocumentStyle.DEFAULT,
        custom_style: Optional[StyleConfig] = None
    ):
        """
        Initialize document builder.

        Args:
            title: Document title
            subtitle: Document subtitle
            author: Document author
            style: Pre-defined style
            custom_style: Custom style configuration
        """
        self.title = title
        self.subtitle = subtitle
        self.author = author
        self.style_config = custom_style or StyleConfig.from_style(style)
        self._docx_available = self._check_docx()
        self._content: List[Dict] = []

        # Add title if provided
        if title:
            self._content.append({
                'type': 'title',
                'text': title,
                'subtitle': subtitle
            })

    def _check_docx(self) -> bool:
        """Check if python-docx is available"""
        try:
            from docx import Document
            return True
        except ImportError:
            return False

    def add_heading(self, text: str, level: int = 1) -> 'DocumentBuilder':
        """
        Add a heading.

        Args:
            text: Heading text
            level: Heading level (1-3)
        """
        self._content.append({
            'type': 'heading',
            'text': text,
            'level': min(max(level, 1), 3)
        })
        return self

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False) -> 'DocumentBuilder':
        """
        Add a paragraph.

        Args:
            text: Paragraph text
            bold: Make text bold
            italic: Make text italic
        """
        self._content.append({
            'type': 'paragraph',
            'text': text,
            'bold': bold,
            'italic': italic
        })
        return self

    def add_bullet_list(self, items: List[str]) -> 'DocumentBuilder':
        """Add a bulleted list"""
        self._content.append({
            'type': 'bullet_list',
            'items': items
        })
        return self

    def add_numbered_list(self, items: List[str]) -> 'DocumentBuilder':
        """Add a numbered list"""
        self._content.append({
            'type': 'numbered_list',
            'items': items
        })
        return self

    def add_table(
        self,
        data: Union[List[Dict], List[List]],
        headers: Optional[List[str]] = None,
        title: str = ""
    ) -> 'DocumentBuilder':
        """
        Add a table.

        Args:
            data: Table data (list of dicts or list of lists)
            headers: Column headers (auto-detected from dict keys if not provided)
            title: Optional table title
        """
        # Normalize data to list of dicts
        if data and isinstance(data[0], list):
            if headers:
                data = [dict(zip(headers, row)) for row in data]
            else:
                headers = [f"Column {i+1}" for i in range(len(data[0]))]
                data = [dict(zip(headers, row)) for row in data]
        elif data and isinstance(data[0], dict) and not headers:
            headers = list(data[0].keys())

        self._content.append({
            'type': 'table',
            'data': data,
            'headers': headers or [],
            'title': title
        })
        return self

    def add_key_value_table(self, data: Dict[str, Any], title: str = "") -> 'DocumentBuilder':
        """
        Add a two-column key-value table.

        Args:
            data: Dictionary of key-value pairs
            title: Optional table title
        """
        table_data = [{'Key': k, 'Value': v} for k, v in data.items()]
        return self.add_table(table_data, headers=['Key', 'Value'], title=title)

    def add_horizontal_rule(self) -> 'DocumentBuilder':
        """Add a horizontal rule/divider"""
        self._content.append({'type': 'horizontal_rule'})
        return self

    def add_page_break(self) -> 'DocumentBuilder':
        """Add a page break"""
        self._content.append({'type': 'page_break'})
        return self

    def add_image(self, path: Union[str, Path], width: Optional[float] = None, caption: str = "") -> 'DocumentBuilder':
        """
        Add an image.

        Args:
            path: Path to image file
            width: Width in inches (auto if None)
            caption: Image caption
        """
        self._content.append({
            'type': 'image',
            'path': str(path),
            'width': width,
            'caption': caption
        })
        return self

    def add_quote(self, text: str, author: str = "") -> 'DocumentBuilder':
        """Add a block quote"""
        self._content.append({
            'type': 'quote',
            'text': text,
            'author': author
        })
        return self

    def save(self, path: Union[str, Path]) -> Path:
        """
        Save document to file.

        Args:
            path: Output file path

        Returns:
            Path to saved file
        """
        path = Path(path)

        if path.suffix.lower() == '.docx' and self._docx_available:
            return self._save_docx(path)
        elif path.suffix.lower() == '.md':
            return self._save_markdown(path)
        elif path.suffix.lower() == '.html':
            return self._save_html(path)
        else:
            # Default to markdown if docx not available
            if not self._docx_available:
                print("Warning: python-docx not available, saving as markdown")
                return self._save_markdown(path.with_suffix('.md'))
            return self._save_docx(path)

    def _save_docx(self, path: Path) -> Path:
        """Save as Word document"""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(self.style_config.margin_top)
            section.bottom_margin = Inches(self.style_config.margin_bottom)
            section.left_margin = Inches(self.style_config.margin_left)
            section.right_margin = Inches(self.style_config.margin_right)

        def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
            """Convert hex color to RGB tuple"""
            hex_color = hex_color.lstrip('#')
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

        for item in self._content:
            item_type = item.get('type')

            if item_type == 'title':
                p = doc.add_paragraph()
                run = p.add_run(item['text'])
                run.bold = True
                run.font.size = Pt(self.style_config.title_size)
                run.font.name = self.style_config.title_font
                rgb = hex_to_rgb(self.style_config.title_color)
                run.font.color.rgb = RGBColor(*rgb)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if item.get('subtitle'):
                    p = doc.add_paragraph()
                    run = p.add_run(item['subtitle'])
                    run.italic = True
                    run.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()  # Spacing

            elif item_type == 'heading':
                level = item.get('level', 1)
                heading = doc.add_heading(item['text'], level=level)

                # Apply custom styling
                for run in heading.runs:
                    if level == 1:
                        run.font.name = self.style_config.heading1_font
                        run.font.size = Pt(self.style_config.heading1_size)
                        rgb = hex_to_rgb(self.style_config.heading1_color)
                    else:
                        run.font.name = self.style_config.heading2_font
                        run.font.size = Pt(self.style_config.heading2_size)
                        rgb = hex_to_rgb(self.style_config.heading2_color)
                    run.font.color.rgb = RGBColor(*rgb)

            elif item_type == 'paragraph':
                p = doc.add_paragraph()
                run = p.add_run(item['text'])
                run.bold = item.get('bold', False)
                run.italic = item.get('italic', False)
                run.font.name = self.style_config.body_font
                run.font.size = Pt(self.style_config.body_size)

            elif item_type == 'bullet_list':
                for list_item in item['items']:
                    p = doc.add_paragraph(list_item, style='List Bullet')
                    for run in p.runs:
                        run.font.name = self.style_config.body_font
                        run.font.size = Pt(self.style_config.body_size)

            elif item_type == 'numbered_list':
                for list_item in item['items']:
                    p = doc.add_paragraph(list_item, style='List Number')
                    for run in p.runs:
                        run.font.name = self.style_config.body_font
                        run.font.size = Pt(self.style_config.body_size)

            elif item_type == 'table':
                if item.get('title'):
                    p = doc.add_paragraph()
                    run = p.add_run(item['title'])
                    run.bold = True

                headers = item['headers']
                data = item['data']

                if headers and data:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'

                    # Header row
                    header_row = table.rows[0]
                    for i, header in enumerate(headers):
                        cell = header_row.cells[i]
                        cell.text = str(header)
                        # Style header
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), self.style_config.table_header_color)
                        cell._tc.get_or_add_tcPr().append(shading)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True
                                rgb = hex_to_rgb(self.style_config.table_header_text_color)
                                run.font.color.rgb = RGBColor(*rgb)

                    # Data rows
                    for row_data in data:
                        row = table.add_row()
                        for i, header in enumerate(headers):
                            cell = row.cells[i]
                            cell.text = str(row_data.get(header, ''))

                doc.add_paragraph()  # Spacing

            elif item_type == 'horizontal_rule':
                p = doc.add_paragraph()
                p.add_run('_' * 50)

            elif item_type == 'page_break':
                doc.add_page_break()

            elif item_type == 'image':
                try:
                    width = Inches(item['width']) if item.get('width') else Inches(5)
                    doc.add_picture(item['path'], width=width)
                    if item.get('caption'):
                        p = doc.add_paragraph()
                        run = p.add_run(item['caption'])
                        run.italic = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    p = doc.add_paragraph()
                    p.add_run(f"[Image: {item['path']}]")

            elif item_type == 'quote':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(f'"{item["text"]}"')
                run.italic = True
                if item.get('author'):
                    p.add_run(f"\n— {item['author']}")

        doc.save(path)
        return path

    def _save_markdown(self, path: Path) -> Path:
        """Save as Markdown"""
        lines = []

        for item in self._content:
            item_type = item.get('type')

            if item_type == 'title':
                lines.append(f"# {item['text']}")
                if item.get('subtitle'):
                    lines.append(f"*{item['subtitle']}*")
                lines.append('')

            elif item_type == 'heading':
                level = item.get('level', 1)
                prefix = '#' * (level + 1)  # +1 because title is #
                lines.append(f"{prefix} {item['text']}")
                lines.append('')

            elif item_type == 'paragraph':
                text = item['text']
                if item.get('bold'):
                    text = f"**{text}**"
                if item.get('italic'):
                    text = f"*{text}*"
                lines.append(text)
                lines.append('')

            elif item_type == 'bullet_list':
                for list_item in item['items']:
                    lines.append(f"- {list_item}")
                lines.append('')

            elif item_type == 'numbered_list':
                for i, list_item in enumerate(item['items'], 1):
                    lines.append(f"{i}. {list_item}")
                lines.append('')

            elif item_type == 'table':
                headers = item['headers']
                data = item['data']

                if item.get('title'):
                    lines.append(f"**{item['title']}**")
                    lines.append('')

                if headers:
                    lines.append('| ' + ' | '.join(headers) + ' |')
                    lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
                    for row in data:
                        cells = [str(row.get(h, '')) for h in headers]
                        lines.append('| ' + ' | '.join(cells) + ' |')
                    lines.append('')

            elif item_type == 'horizontal_rule':
                lines.append('---')
                lines.append('')

            elif item_type == 'page_break':
                lines.append('\n---\n')

            elif item_type == 'image':
                caption = item.get('caption', 'Image')
                lines.append(f"![{caption}]({item['path']})")
                lines.append('')

            elif item_type == 'quote':
                lines.append(f"> {item['text']}")
                if item.get('author'):
                    lines.append(f"> — {item['author']}")
                lines.append('')

        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

        return path

    def _save_html(self, path: Path) -> Path:
        """Save as HTML"""
        html_parts = [
            '<!DOCTYPE html>',
            '<html>',
            '<head>',
            f'<title>{self.title}</title>',
            '<style>',
            f'body {{ font-family: {self.style_config.body_font}, sans-serif; margin: 40px; line-height: 1.6; }}',
            f'h1 {{ color: #{self.style_config.title_color}; }}',
            f'h2 {{ color: #{self.style_config.heading1_color}; }}',
            f'h3 {{ color: #{self.style_config.heading2_color}; }}',
            'table { border-collapse: collapse; width: 100%; margin: 20px 0; }',
            f'th {{ background-color: #{self.style_config.table_header_color}; color: #{self.style_config.table_header_text_color}; padding: 12px; text-align: left; }}',
            f'td {{ border: 1px solid #ddd; padding: 12px; }}',
            f'tr:nth-child(even) {{ background-color: #{self.style_config.table_row_alt_color}; }}',
            'blockquote { border-left: 4px solid #ccc; margin: 20px 0; padding-left: 20px; font-style: italic; }',
            'ul, ol { margin: 20px 0; }',
            'li { margin: 5px 0; }',
            '</style>',
            '</head>',
            '<body>'
        ]

        for item in self._content:
            item_type = item.get('type')

            if item_type == 'title':
                html_parts.append(f"<h1>{item['text']}</h1>")
                if item.get('subtitle'):
                    html_parts.append(f"<p><em>{item['subtitle']}</em></p>")

            elif item_type == 'heading':
                level = item.get('level', 1) + 1
                html_parts.append(f"<h{level}>{item['text']}</h{level}>")

            elif item_type == 'paragraph':
                text = item['text']
                if item.get('bold'):
                    text = f"<strong>{text}</strong>"
                if item.get('italic'):
                    text = f"<em>{text}</em>"
                html_parts.append(f"<p>{text}</p>")

            elif item_type == 'bullet_list':
                html_parts.append('<ul>')
                for list_item in item['items']:
                    html_parts.append(f"<li>{list_item}</li>")
                html_parts.append('</ul>')

            elif item_type == 'numbered_list':
                html_parts.append('<ol>')
                for list_item in item['items']:
                    html_parts.append(f"<li>{list_item}</li>")
                html_parts.append('</ol>')

            elif item_type == 'table':
                if item.get('title'):
                    html_parts.append(f"<p><strong>{item['title']}</strong></p>")

                html_parts.append('<table>')
                headers = item['headers']
                if headers:
                    html_parts.append('<tr>' + ''.join(f'<th>{h}</th>' for h in headers) + '</tr>')
                for row in item['data']:
                    cells = ''.join(f'<td>{row.get(h, "")}</td>' for h in headers)
                    html_parts.append(f'<tr>{cells}</tr>')
                html_parts.append('</table>')

            elif item_type == 'horizontal_rule':
                html_parts.append('<hr>')

            elif item_type == 'page_break':
                html_parts.append('<div style="page-break-after: always;"></div>')

            elif item_type == 'image':
                width = f'width="{item["width"] * 96}px"' if item.get('width') else ''
                html_parts.append(f'<img src="{item["path"]}" {width} alt="{item.get("caption", "")}">')
                if item.get('caption'):
                    html_parts.append(f'<p><em>{item["caption"]}</em></p>')

            elif item_type == 'quote':
                html_parts.append(f'<blockquote>{item["text"]}')
                if item.get('author'):
                    html_parts.append(f'<br>— {item["author"]}')
                html_parts.append('</blockquote>')

        html_parts.extend(['</body>', '</html>'])

        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_parts))

        return path


class OnePagerBuilder(DocumentBuilder):
    """
    Specialized builder for one-page executive summaries.

    Creates compact, professional one-pagers with tight formatting.
    """

    def __init__(
        self,
        title: str,
        subtitle: str = "",
        style: DocumentStyle = DocumentStyle.EXECUTIVE
    ):
        super().__init__(
            title=title,
            subtitle=subtitle,
            style=style
        )
        # Tighter margins for one-pagers
        self.style_config.margin_top = 0.5
        self.style_config.margin_bottom = 0.5
        self.style_config.margin_left = 0.75
        self.style_config.margin_right = 0.75
        self.style_config.body_size = 10


# ==================== Factory Functions ====================


def create_simple_document(title: str, content: str) -> DocumentBuilder:
    """Create a simple document with title and content"""
    doc = DocumentBuilder(title=title)
    doc.add_paragraph(content)
    return doc


def create_report_document(
    title: str,
    sections: Dict[str, str],
    style: DocumentStyle = DocumentStyle.CORPORATE
) -> DocumentBuilder:
    """
    Create a report with multiple sections.

    Args:
        title: Document title
        sections: Dict of section_title -> content
        style: Document style
    """
    doc = DocumentBuilder(title=title, style=style)
    for section_title, content in sections.items():
        doc.add_heading(section_title, level=1)
        doc.add_paragraph(content)
    return doc


def create_proposal_document(
    title: str,
    executive_summary: str,
    problem_statement: str,
    proposed_solution: str,
    benefits: List[str],
    timeline: Optional[List[Dict]] = None,
    pricing: Optional[Dict[str, Any]] = None
) -> DocumentBuilder:
    """Create a proposal document with standard sections"""
    doc = DocumentBuilder(title=title, style=DocumentStyle.CORPORATE)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(executive_summary)

    doc.add_heading("Problem Statement", level=1)
    doc.add_paragraph(problem_statement)

    doc.add_heading("Proposed Solution", level=1)
    doc.add_paragraph(proposed_solution)

    doc.add_heading("Key Benefits", level=1)
    doc.add_bullet_list(benefits)

    if timeline:
        doc.add_heading("Timeline", level=1)
        doc.add_table(timeline, headers=['Phase', 'Description', 'Duration'])

    if pricing:
        doc.add_heading("Investment", level=1)
        doc.add_key_value_table(pricing)

    return doc


# ==================== Exports ====================

__all__ = [
    'DocumentBuilder',
    'DocumentStyle',
    'StyleConfig',
    'OnePagerBuilder',
    'create_simple_document',
    'create_report_document',
    'create_proposal_document',
]
